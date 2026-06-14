from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def generate_pdf_report(review_data, output_path):

    doc = SimpleDocTemplate(output_path)

    styles = getSampleStyleSheet()

    elements = [

        Paragraph(
            "TestLens AI Review Report",
            styles["Title"]
        ),

        Paragraph(
            f"Overall Score: {review_data.get('overall_score',0)}",
            styles["Normal"]
        ),

        Paragraph(
            f"Coverage: {review_data.get('coverage_percentage',0)}%",
            styles["Normal"]
        ),

        Paragraph(
            f"Risk Level: {review_data.get('risk_level','Low')}",
            styles["Normal"]
        )

    ]

    doc.build(elements)

    return output_path

def generate_docx_report(review_data, output_path):

    doc = Document()

    doc.add_heading(
        "TestLens AI Review Report",
        0
    )

    doc.add_paragraph(
        f"Overall Score: {review_data.get('overall_score',0)}"
    )

    doc.add_paragraph(
        f"Coverage: {review_data.get('coverage_percentage',0)}%"
    )

    doc.add_paragraph(
        f"Risk Level: {review_data.get('risk_level','Low')}"
    )

    doc.save(output_path)

    return output_path